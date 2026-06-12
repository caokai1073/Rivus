"""
ingest.py — Content parsing + chunking + embedding + database ingestion
"""
import re
from pathlib import Path
from typing import Optional

import requests
from readability import Document  # readability-lxml
from sentence_transformers import SentenceTransformer

import threading
from db import init_db, insert_document, insert_chunk, update_summary

# ── Global singleton embedding model ─────────────────────────────────────────
_model: Optional[SentenceTransformer] = None

def get_embed_model() -> SentenceTransformer:
    global _model
    if _model is None:
        print("[ingest] Loading embedding model BAAI/bge-m3 ...")
        # Prefer offline loading (no network if model is already cached)
        try:
            _model = SentenceTransformer("BAAI/bge-m3", local_files_only=True)
        except Exception:
            # First run: download the model
            _model = SentenceTransformer("BAAI/bge-m3")
    return _model


# ── Text chunking (paragraph-aware) ──────────────────────────────────────────
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

def chunk_text(text: str) -> list[str]:
    """
    Paragraph-aware chunking: split at paragraph boundaries to preserve semantic coherence.
    Each chunk is at most CHUNK_SIZE characters; adjacent chunks overlap by CHUNK_OVERLAP characters.
    """
    # Normalize excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    # Split on paragraph boundaries (double newline)
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]

    chunks = []
    current = ""

    for para in paragraphs:
        # Force-split paragraphs that exceed CHUNK_SIZE
        if len(para) > CHUNK_SIZE:
            if current:
                chunks.append(current.strip())
                current = ""
            # Try to split on sentence boundaries (。！？.!?)
            sentences = re.split(r"(?<=[。！？.!?])\s*", para)
            buf = ""
            for sent in sentences:
                if len(buf) + len(sent) > CHUNK_SIZE and buf:
                    chunks.append(buf.strip())
                    # Overlap: keep last CHUNK_OVERLAP characters
                    buf = buf[-CHUNK_OVERLAP:] + sent
                else:
                    buf += sent
            if buf.strip():
                chunks.append(buf.strip())
            continue

        # Adding paragraph would exceed limit → save current chunk, start new one (with overlap)
        if current and len(current) + len(para) + 2 > CHUNK_SIZE:
            chunks.append(current.strip())
            # Overlap: carry over trailing content from previous chunk
            overlap = current[-CHUNK_OVERLAP:].strip()
            current = overlap + "\n\n" + para if overlap else para
        else:
            current = current + "\n\n" + para if current else para

    if current.strip():
        chunks.append(current.strip())

    return [c for c in chunks if c]


# ── URL parsing ───────────────────────────────────────────────────────────────
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}

def parse_url(url: str) -> tuple[str, str, str]:
    resp = requests.get(url, headers=HEADERS, timeout=15)
    resp.raise_for_status()
    doc = Document(resp.text)
    title = doc.title() or url
    html_content = doc.summary()
    plain = re.sub(r"<[^>]+>", " ", html_content)
    plain = re.sub(r"[ \t]+", " ", plain)
    plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
    from urllib.parse import urlparse
    source = urlparse(url).netloc or url
    return title, plain, source


# ── PDF parsing ───────────────────────────────────────────────────────────────

def _clean_pdf_text(text: str) -> str:
    """Clean common artifacts and formatting issues from extracted PDF text"""
    # Merge hyphenated line breaks in English words
    text = re.sub(r"-\n([a-z])", r"\1", text)
    # Single newlines become spaces (intra-paragraph); double newlines preserved (paragraphs)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    # Collapse multiple spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf(file_path: str) -> tuple[str, str, str]:
    """
    Parse a PDF; prefer pymupdf (better quality), fall back to pypdf.
    Returns (title, text, source).
    """
    # ── Option 1: pymupdf ─────────────────────────────────────────────────────
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)

        # Use metadata title if available, otherwise use filename
        meta_title = (doc.metadata or {}).get("title", "").strip()
        title = meta_title or Path(file_path).stem

        pages = []
        for page in doc:
            # sort=True orders text blocks by reading order (better for multi-column PDFs)
            page_text = page.get_text("text", sort=True)
            if page_text.strip():
                pages.append(page_text)

        full_text = _clean_pdf_text("\n\n".join(pages))

        if len(full_text) < 50:
            raise ValueError("Scanned PDF (image-only); cannot extract text. Please convert with an OCR tool first.")

        return title, full_text, "PDF"

    except ImportError:
        pass  # pymupdf not installed, fall back to pypdf

    # ── Option 2: pypdf (fallback) ────────────────────────────────────────────
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        meta = reader.metadata or {}
        title = (meta.get("/Title") or "").strip() or Path(file_path).stem
        pages = [page.extract_text() or "" for page in reader.pages]
        full_text = _clean_pdf_text("\n\n".join(pages))
        if not full_text:
            raise ValueError("PDF content is empty; it may be a scanned or encrypted file.")
        return title, full_text, "PDF"
    except ImportError:
        raise ImportError("Please install pymupdf: pip install pymupdf")


# ── Markdown parsing ──────────────────────────────────────────────────────────

def parse_md(file_path: str) -> tuple[str, str, str]:
    """
    Parse a Markdown file, returns (title, plain_text, source).
    - Title: first H1 heading, or filename if none found
    - Strips markdown syntax to produce plain text for embedding
    """
    text = Path(file_path).read_text(encoding="utf-8", errors="ignore")

    # Extract title from first H1
    title = ""
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            break
    if not title:
        title = Path(file_path).stem

    # Strip markdown syntax
    clean = text
    # Fenced code blocks: remove fence, keep content
    clean = re.sub(r"```[^\n]*\n([\s\S]*?)```", r"\1", clean)
    # Inline code
    clean = re.sub(r"`([^`\n]+)`", r"\1", clean)
    # Headings
    clean = re.sub(r"^#{1,6}\s+", "", clean, flags=re.MULTILINE)
    # Bold / italic / bold-italic
    clean = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", clean)
    clean = re.sub(r"_{1,3}([^_\n]+)_{1,3}", r"\1", clean)
    # Links [text](url) → text
    clean = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", clean)
    # Images ![alt](url) → alt
    clean = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", clean)
    # Blockquote markers
    clean = re.sub(r"^>\s*", "", clean, flags=re.MULTILINE)
    # Horizontal rules
    clean = re.sub(r"^[-*_]{3,}\s*$", "", clean, flags=re.MULTILINE)
    # Unordered list markers
    clean = re.sub(r"^[ \t]*[-*+]\s+", "", clean, flags=re.MULTILINE)
    # Ordered list markers
    clean = re.sub(r"^[ \t]*\d+\.\s+", "", clean, flags=re.MULTILINE)
    # Collapse excessive blank lines
    clean = re.sub(r"\n{3,}", "\n\n", clean)

    return title, clean.strip(), "Markdown"


# ── Excel parsing ─────────────────────────────────────────────────────────────

def parse_excel(file_path: str) -> tuple[str, str, str]:
    """
    Parse an Excel (.xlsx/.xls) file, returns (title, plain_text, source).
    Each sheet is converted to plain text (tab-separated); sheet name used as section header.
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is not installed. Run: pip install openpyxl")

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    title = Path(file_path).stem

    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Skip fully empty rows
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))

    wb.close()

    if not parts:
        raise ValueError("Excel file is empty. Please check that the file is valid.")

    full_text = "\n\n".join(parts)
    return title, full_text, "Excel"


# ── Main entry points ─────────────────────────────────────────────────────────

def ingest_url(url: str, custom_title: str = None, folder_id: int = None, progress_cb=None) -> dict:
    init_db()
    print(f"[ingest] Fetching: {url}")
    title, text, source = parse_url(url)
    if custom_title:
        title = custom_title
    return _ingest_text(url=url, title=title, text=text, source=source, folder_id=folder_id, progress_cb=progress_cb)


def ingest_text(title: str, text: str, source: str = "User", folder_id: int = None, progress_cb=None) -> dict:
    init_db()
    return _ingest_text(url="", title=title, text=text, source=source, folder_id=folder_id, progress_cb=progress_cb)


def ingest_docx(file_path: str, custom_title: str = None, folder_id: int = None, stored_url: str = "", progress_cb=None) -> dict:
    """Parse a Word (.docx) document"""
    init_db()
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    title = custom_title or Path(file_path).stem

    # Extract paragraphs + table text
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))

    full_text = "\n\n".join(parts)
    if not full_text.strip():
        raise ValueError("Word document is empty. Please check that the file is valid.")

    return _ingest_text(url=stored_url or file_path, title=title, text=full_text, source="Word", folder_id=folder_id, progress_cb=progress_cb)


def ingest_pdf(file_path: str, custom_title: str = None, folder_id: int = None, stored_url: str = "", progress_cb=None) -> dict:
    init_db()
    title, text, source = parse_pdf(file_path)
    if custom_title:
        title = custom_title
    return _ingest_text(url=stored_url, title=title, text=text, source=source, folder_id=folder_id, progress_cb=progress_cb)


def ingest_excel(file_path: str, custom_title: str = None, folder_id: int = None, stored_url: str = "", progress_cb=None) -> dict:
    """Parse an Excel (.xlsx) file"""
    init_db()
    title, text, source = parse_excel(file_path)
    if custom_title:
        title = custom_title
    elif stored_url:
        # Use "{timestamp}.xlsx" — unique and includes extension so AI can identify file type
        p = Path(stored_url)
        title = p.stem.split('_')[0] + p.suffix
    return _ingest_text(url=stored_url or file_path, title=title, text=text, source=source, folder_id=folder_id, progress_cb=progress_cb)


def ingest_md(file_path: str, custom_title: str = None, folder_id: int = None, stored_url: str = "", progress_cb=None) -> dict:
    """Parse a Markdown (.md) file"""
    init_db()
    title, text, source = parse_md(file_path)
    if custom_title:
        title = custom_title
    elif stored_url and title == Path(file_path).stem:
        # parse_md fell back to the tmp filename; use "{timestamp}.md" instead
        p = Path(stored_url)
        title = p.stem.split('_')[0] + p.suffix
    if not text.strip():
        raise ValueError("Markdown file is empty. Please check that the file is valid.")
    return _ingest_text(url=stored_url or file_path, title=title, text=text, source=source, folder_id=folder_id, progress_cb=progress_cb)


_enrich_semaphore = threading.Semaphore(1)  # allow only one summary generation task at a time

def _ingest_text(url: str, title: str, text: str, source: str,
                 folder_id: int = None, progress_cb=None) -> dict:
    if not text.strip():
        raise ValueError("Extracted content is empty. Please check the link or input.")

    chunks = chunk_text(text)
    total = len(chunks)
    print(f"[ingest] '{title}' split into {total} chunks")
    if progress_cb:
        progress_cb(0, total, "chunked")

    # Batch encode to avoid OOM/CPU overload on large documents
    model = get_embed_model()
    BATCH = 16
    all_embeddings = []
    done = 0
    for i in range(0, total, BATCH):
        batch = chunks[i:i + BATCH]
        embs = model.encode(batch, show_progress_bar=False,
                            normalize_embeddings=True, batch_size=BATCH)
        all_embeddings.extend(embs.tolist())
        done += len(batch)
        if progress_cb:
            progress_cb(done, total, "embedding")

    if progress_cb:
        progress_cb(total, total, "saving")

    doc_id = insert_document(url=url, title=title, source=source, full_text=text, folder_id=folder_id)
    for i, (chunk, emb) in enumerate(zip(chunks, all_embeddings)):
        insert_chunk(doc_id=doc_id, chunk_index=i, text=chunk, embedding=emb)

    print(f"[ingest] OK Saved, doc_id={doc_id}")

    # Generate summary in background (semaphore ensures only one runs at a time)
    def _enrich_guarded(doc_id, title, text):
        with _enrich_semaphore:
            _enrich_background(doc_id, title, text)

    threading.Thread(
        target=_enrich_guarded,
        args=(doc_id, title, text),
        daemon=True
    ).start()

    return {
        "doc_id": doc_id,
        "title": title,
        "source": source,
        "chunks": total,
        "chars": len(text),
    }


def _enrich_background(doc_id: int, title: str, text: str):
    """Background task: call Ollama to generate a document summary and key info"""
    try:
        import requests as req
        from query import DEFAULT_MODEL

        # Use first 4000 chars (covers core content while keeping inference fast)
        excerpt = text[:4000]
        prompt = f"""Analyze the following document and extract key information.

Document title: {title}

Document content:
{excerpt}

Output the following (reply in the same language as the document):
[Summary] 2-3 sentences summarizing the core content
[Key Points] 3-5 most important points (one per line, starting with "-")
[Keywords] 3-6 keywords, comma-separated"""

        resp = req.post(
            "http://localhost:11434/api/chat",
            json={
                "model": DEFAULT_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
            },
            timeout=120,
        )
        if resp.ok:
            summary = resp.json()["message"]["content"].strip()
            update_summary(doc_id, summary)
            print(f"[enrich] OK doc_id={doc_id} summary generated")
        else:
            print(f"[enrich] Ollama returned error: {resp.status_code}")
    except Exception as e:
        print(f"[enrich] Summary generation failed doc_id={doc_id}: {e}")
